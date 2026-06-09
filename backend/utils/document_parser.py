from pathlib import Path

from pypdf import PdfReader
from docx import Document


def parse_txt_file(file_path: Path) -> str:
    """
    解析 txt / md 等纯文本文件
    """
    try:
        return file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="gbk", errors="ignore")


def parse_pdf_file(file_path: Path) -> str:
    """
    解析文字型 PDF。
    注意：扫描版 PDF 无法直接提取文字，需要 OCR。
    """
    reader = PdfReader(str(file_path))

    texts = []

    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""

        if text.strip():
            texts.append(f"第 {page_index + 1} 页：\n{text}")

    return "\n\n".join(texts)


def parse_docx_file(file_path: Path) -> str:
    """
    解析 docx Word 文档
    """
    doc = Document(str(file_path))

    texts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                texts.append(" | ".join(row_text))

    return "\n".join(texts)


def parse_document(file_path: Path) -> str:
    """
    根据文件后缀自动选择解析方式
    """
    suffix = file_path.suffix.lower()

    if suffix in [".txt", ".md"]:
        text = parse_txt_file(file_path)

    elif suffix == ".pdf":
        text = parse_pdf_file(file_path)

    elif suffix == ".docx":
        text = parse_docx_file(file_path)

    else:
        raise ValueError(
            f"暂不支持的文件类型：{suffix}，目前支持 .txt / .md / .pdf / .docx"
        )

    text = text.strip()

    if not text:
        raise ValueError(
            "文档内容为空，可能是扫描版 PDF，或者文件中没有可提取的文本。"
        )

    return text